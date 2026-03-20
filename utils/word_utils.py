# ============================================================
# IMPORTS - word_utils.py
# ============================================================

from io import BytesIO
from docx.shared import Inches
import logging
log = logging.getLogger(__name__)
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC

def add_formatted_section_to_word(driver, wait, document, section_type):
    """
    Handles formatted Word sections with screenshots.
    Supports:
        - daily_activity
        - payment_dashboard
        - generic (fallback)

    Screenshots are inserted directly from memory.
    No PNG files are saved on disk.
    """

    document.add_paragraph("")

    # Helper → insert screenshot directly from memory
    def insert_screenshot():
        screenshot_bytes = driver.get_screenshot_as_png()
        image_stream = BytesIO(screenshot_bytes)
        document.add_picture(image_stream, width=Inches(6))


    # =========================================================
    # STEP 2 - DAILY ACTIVITY (2 screenshots)
    # =========================================================
    if section_type == "daily_activity":

        log.info("📸 Adding Step 2 - Daily Activity section...")

        document.add_heading("Step 2 Daily Activity Report", level=2)
        document.add_paragraph(
            "Move out/ Move In have been scheduled on AME date."
        )

        # -----------------------------
        # Screenshot 1 - Dashboard
        # -----------------------------
        insert_screenshot()
        log.info("✅ Dashboard screenshot added.")

        # -----------------------------
        # Click Deposit Accounting
        # -----------------------------
        try:
            driver.switch_to.default_content()

            # Switch to correct iframe
            wait.until(EC.frame_to_be_available_and_switch_to_it((By.ID, "filter")))

            deposit_link = wait.until(
                EC.element_to_be_clickable((By.ID, "DepAcctgLink"))
            )

            driver.execute_script("arguments[0].click();", deposit_link)
            log.info("🔎 Deposit Accounting clicked. Waiting for load...")
            time.sleep(4)

            # -----------------------------
            # Screenshot 2 - Deposit Accounting
            # -----------------------------
            insert_screenshot()
            log.info("✅ Deposit Accounting screenshot added.")

        except Exception as e:
            log.warning(f"⚠️ Could not capture Deposit Accounting screenshot: {e}")

        finally:
            driver.switch_to.default_content()

    # =========================================================
    # STEP 4 - Credit Card DASHBOARD
    # =========================================================
    elif section_type == "credit_card_dashboard":

        log.info("📸 Adding Step 4 - Yardi Credit Card Dashboard section...")

        document.add_heading("Step 4 Yardi Credit Card Dashboard", level=2)
        document.add_paragraph(
            "Below are the pending points for Credit card dashboard."
        )

        screenshot_bytes = driver.get_screenshot_as_png()
        image_stream = BytesIO(screenshot_bytes)

        document.add_picture(image_stream, width=Inches(6))

        log.info("✅ Credit Card Dashboard screenshot added.")



    # =========================================================
    # STEP 5 - PAYMENT DASHBOARD
    # =========================================================
    elif section_type == "payment_dashboard":

        log.info("📸 Adding Step 5 - Payment Dashboard section...")

        document.add_heading("Step 5 Payment Dashboard", level=2)
        document.add_paragraph(
            "Below are the pending points for Payment dashboard."
        )

        insert_screenshot()
        log.info("✅ Payment Dashboard screenshot added.")

    # =========================================================
    # STEP 7 - Open Batches
    # =========================================================


    elif section_type == "transaction_register":

        log.info("📸 Adding Step 7 - Clear Open Batches section...")

      
        document.add_heading("Step 7 Clear Open Batches", level=2)
        document.add_paragraph(
            "There are some open batches."
        )
        screenshot_bytes = driver.get_screenshot_as_png()
        image_stream = BytesIO(screenshot_bytes)

        document.add_picture(image_stream, width=Inches(6))

        log.info("✅ Clear Open Batches screenshot added to Word.")


        
    # =========================================================
    # STEP 8 - NSF
    # =========================================================



    elif section_type == "nsf_payments":

        log.info("📸 Adding Step 8 - NSF Payments section...")

        document.add_heading("Step 8 NSF Payments", level=2)

        document.add_paragraph(
            "Below are the Negative amounts showing in NSF payments."
        )

        screenshot_bytes = driver.get_screenshot_as_png()
        image_stream = BytesIO(screenshot_bytes)

        document.add_picture(image_stream, width=Inches(6))

        log.info("✅ NSF Payments screenshot added to Word.")

        
    # =========================================================
    # STEP 9 - Conver to Prepay
    # =========================================================


    elif section_type == "apply_open_credits":

        document.add_heading("Step 9 Apply Open Credits", level=2)
        document.add_paragraph(
            "Below open credits are there in the reports."
        )

        screenshot_bytes = driver.get_screenshot_as_png()
        image_stream = BytesIO(screenshot_bytes)
        document.add_picture(image_stream, width=Inches(6))

        log.info("✅ Apply Open Credits screenshot added.")
        
    # =========================================================
    # STEP 10 Write Offs
    # =========================================================

    elif section_type == "write_offs":

        log.info("📸 Adding Step 10 - Write Offs section...")

        document.add_heading("Step 10 Write Offs", level=2)
        document.add_paragraph(
            "Below Past/Canceled, Denied resident are there in the report."
        )

        screenshot_bytes = driver.get_screenshot_as_png()
        image_stream = BytesIO(screenshot_bytes)

        document.add_picture(image_stream, width=Inches(6))

        log.info("✅ Write Offs screenshot added.")

    # =========================================================
    # STEP 13 Payable 
    # =========================================================

    elif section_type == "payable_aging":

        log.info("📸 Adding Step 13 - Verify Unclaimed Refund Process section...")

        document.add_heading(
            "Step 13 Verify Unclaimed Refund Process",
            level=2
        )

        document.add_paragraph(
            "There are some payable over due than 90 days."
        )

        screenshot_bytes = driver.get_screenshot_as_png()
        image_stream = BytesIO(screenshot_bytes)

        document.add_picture(image_stream, width=Inches(6))

        log.info("✅ Step 13 screenshot added.")



    # =========================================================
    # GENERIC SNAPSHOT
    # =========================================================
    else:

        log.info("📸 Adding generic screenshot section...")

        document.add_heading("Observation Snapshot", level=2)

        insert_screenshot()
        log.info("✅ Generic screenshot added.")

