import os
import time
import logging
import win32com.client as win32
import logging
log = logging.getLogger(__name__)
from io import BytesIO
from PIL import ImageGrab
from datetime import datetime
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support import expected_conditions as EC

from docx.shared import Inches

from utils.download_utils import wait_for_download_and_rename

def run_gpr_activity(driver, wait, prop_code, prop_name, folder_path):

    driver.switch_to.default_content()
    log.info("🔍 Searching for Residential Analytics...")

    wait.until(EC.element_to_be_clickable((By.ID, "miSearch"))).click()

    search_box = wait.until(
        EC.visibility_of_element_located(
            (By.XPATH, "//li[@id='miSearch']//input[@type='text']")
        )
    )
    search_box.clear()
    search_box.send_keys("Residential Analytics")

    wait.until(
        EC.element_to_be_clickable(
            (By.XPATH, "//a[contains(text(),'Residential Analytics')]")
        )
    ).click()

    log.info("📊 Residential Analytics opened.")

    driver.switch_to.default_content()

    for frame in driver.find_elements(By.TAG_NAME, "iframe"):
        driver.switch_to.default_content()
        driver.switch_to.frame(frame)
        if driver.find_elements(By.ID, "ReportType_DropDownList"):
            break
    else:
        raise RuntimeError("❌ Residential Analytics iframe not found.")

    # Set property
    prop_box = wait.until(
        EC.presence_of_element_located((By.ID, "PropLookup_LookupCode"))
    )
    driver.execute_script("arguments[0].value = arguments[1];", prop_box, prop_code)
    prop_box.send_keys(Keys.TAB)

    Select(wait.until(
        EC.presence_of_element_located((By.ID, "ReportType_DropDownList"))
    )).select_by_visible_text("Gross Potential Rent")

    from datetime import datetime
    today = datetime.today().strftime("%m/%d/%Y")

    date_box = wait.until(
        EC.presence_of_element_located((By.ID, "Date2_TextBox"))
    )
    date_box.clear()
    date_box.send_keys(today)
    date_box.send_keys(Keys.TAB)
    log.info(f"📅 Date set to {today}")

    wait.until(EC.element_to_be_clickable((By.ID, "Display_Button"))).click()
    time.sleep(4)

    download_dir = os.path.join(os.path.expanduser("~"), "Downloads")

    # 🔥 IMPORTANT — capture files BEFORE download
    files_before = set(os.listdir(download_dir))

    wait.until(EC.element_to_be_clickable((By.ID, "Excel_Button"))).click()

    log.info("📥 Waiting for GPR Excel download...")

    return wait_for_download_and_rename(
        download_dir=download_dir,
        files_before=files_before,
        target_folder=folder_path,
        base_filename=f"{prop_name}_GPR"
    )


from io import BytesIO
from PIL import ImageGrab
from docx.shared import Inches
import win32com.client as win32
import os
import logging

log = logging.getLogger(__name__)


def review_gpr_report(file_path, document):

    # ---------------------------------------------------------
    # Safety Check
    # ---------------------------------------------------------
    if not os.path.exists(file_path):
        raise Exception(f"GPR file not found: {file_path}")

    excel = win32.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    excel.ScreenUpdating = False
    excel.EnableEvents = False

    wb = excel.Workbooks.Open(file_path)
    ws = wb.ActiveSheet

    last_row = ws.Cells(ws.Rows.Count, 2).End(-4162).Row

    document.add_heading("Gross Potential Rent Review", level=2)

    # =========================================================
    # SAFE NUMBER CONVERTER
    # =========================================================

    def safe_number(value):
        try:
            if value in (None, ""):
                return 0
            return float(str(value).replace(",", "").strip())
        except:
            return 0

    # =========================================================
    # RESET TABLE
    # =========================================================

    def reset_table():
        try:
            if ws.AutoFilterMode:
                ws.AutoFilterMode = False
        except:
            pass

        ws.Range(f"A7:Q{last_row}").Interior.ColorIndex = 0

    # =========================================================
    # COPY ONLY HIGHLIGHTED ROWS
    # =========================================================

    def copy_visible(color_code):

        xlFilterCellColor = 8

        # Excel must repaint to capture picture
        excel.ScreenUpdating = True

        ws.Range(f"A6:Q{last_row}").AutoFilter(
            Field=1,
            Criteria1=color_code,
            Operator=xlFilterCellColor
        )

        time.sleep(0.8)

        rng = ws.Range(f"A6:Q{last_row}")
        rng.CopyPicture(Appearance=1, Format=2)

        time.sleep(0.8)

        img = ImageGrab.grabclipboard()

        if img:
            image_stream = BytesIO()
            img.save(image_stream, format="PNG")
            image_stream.seek(0)
            document.add_picture(image_stream, width=Inches(6))
        else:
            document.add_paragraph("⚠ Unable to capture snapshot.")

        reset_table()

        # Disable again for performance
        excel.ScreenUpdating = False

        
    # =========================================================
    # LOAD DATA ONCE (BIG SPEED BOOST)
    # =========================================================

    data = ws.Range(f"A7:Q{last_row}").Value

    # =========================================================
    # RULE 1 — Market - Actual - Vacancy = Loss/Gain
    # =========================================================

    reset_table()

    color1 = 65535
    highlight = False
    failed_units = []

    for idx, row in enumerate(data, start=7):

        if not row[4]:
            continue

        unit_number = str(row[1] or "").strip()

        market = safe_number(row[5])
        actual = safe_number(row[9])
        vacancy = safe_number(row[8])
        loss_gain = safe_number(row[6])

        calc = round(market - actual - vacancy, 2)

        if round(loss_gain, 2) != calc:

            ws.Range(f"A{idx}:Q{idx}").Interior.Color = color1
            highlight = True
            failed_units.append(unit_number)

    if highlight:

        document.add_heading("GPR – Market - Actual - Vacancy Loss / Gain", level=3)

        document.add_paragraph(
            "Loss / Gain is higher than Market Rent for Unit(s): "
            + ", ".join(failed_units)
        )

        copy_visible(color1)

    else:

        document.add_paragraph("All Looks Good for Loss / Gain to Lease")

    reset_table()

    # =========================================================
    # RULE 2 — Potential - Actual = Vacancy
    # =========================================================

    color2 = 15773696
    highlight = False
    failed_units = []

    for idx, row in enumerate(data, start=7):

        if not row[4]:
            continue

        unit_number = str(row[1] or "").strip()

        potential = safe_number(row[7])
        actual = safe_number(row[9])
        vacancy = safe_number(row[8])

        if round(potential - actual, 2) != round(vacancy, 2):

            ws.Range(f"A{idx}:Q{idx}").Interior.Color = color2
            highlight = True
            failed_units.append(unit_number)

    if highlight:

        document.add_heading("GPR – Potential - Actual Vacancy Loss", level=3)

        document.add_paragraph(
            "Difference of Potential Rent and Actual Rent is not equal to Vacancy Loss for Unit(s): "
            + ", ".join(failed_units)
        )

        copy_visible(color2)

    else:

        document.add_paragraph("All Looks Good for Vacancy Loss")

    reset_table()

    # =========================================================
    # RULE 3 — Actual Rent >= 0
    # =========================================================

    color3 = 255
    highlight = False
    failed_units = []

    for idx, row in enumerate(data, start=7):

        if not row[4]:
            continue

        unit_number = str(row[1] or "").strip()

        actual = safe_number(row[9])

        if actual < 0:

            ws.Range(f"A{idx}:Q{idx}").Interior.Color = color3
            highlight = True
            failed_units.append(unit_number)

    if highlight:

        document.add_heading("GPR – Negative Actual Rent Detected", level=3)

        document.add_paragraph(
            "Actual Rent is Negative for Unit(s): "
            + ", ".join(failed_units)
        )

        copy_visible(color3)

    else:

        document.add_paragraph("All Looks Good for Actual Rent")

    reset_table()

    # =========================================================
    # RULE 4 — Concession Check
    # =========================================================

    color4 = 10092543
    highlight = False
    failed_units = []

    for idx, row in enumerate(data, start=7):

        if not row[4]:
            continue

        unit_number = str(row[1] or "").strip()

        concession = safe_number(row[10])
        actual = safe_number(row[9])

        if not (concession <= 0 and actual >= abs(concession)):

            ws.Range(f"A{idx}:Q{idx}").Interior.Color = color4
            highlight = True
            failed_units.append(unit_number)

    if highlight:

        document.add_heading("GPR – Concession Validation Failed", level=3)

        document.add_paragraph(
            "Concession validation failed for Unit(s): "
            + ", ".join(failed_units)
        )

        copy_visible(color4)

    else:

        document.add_paragraph("All Looks Good for Concession")

    reset_table()

    # =========================================================
    # RULE 5 — Write Off <= Actual Rent
    # =========================================================

    color5 = 49407
    highlight = False
    failed_units = []

    for idx, row in enumerate(data, start=7):

        if not row[4]:
            continue

        unit_number = str(row[1] or "").strip()

        writeoff = safe_number(row[11])
        actual = safe_number(row[9])

        if writeoff > actual:

            ws.Range(f"A{idx}:Q{idx}").Interior.Color = color5
            highlight = True
            failed_units.append(unit_number)

    if highlight:

        document.add_heading("GPR – Write Off Greater Than Actual Rent", level=3)

        document.add_paragraph(
            "Write Off is Greater than One Month Actual Rent for Unit(s): "
            + ", ".join(failed_units)
        )

        copy_visible(color5)

    else:

        document.add_paragraph("All Looks Good for Write Off")

    reset_table()

    wb.Close(False)
    excel.Quit()

def run_gpr_full(driver, wait, prop_code, prop_name, folder_path, document):

    gpr_path = run_gpr_activity(
        driver, wait, prop_code, prop_name, folder_path
    )

    review_gpr_report(gpr_path, document)

