import time
from datetime import datetime
from io import BytesIO
import logging
log = logging.getLogger(__name__)
from selenium.webdriver.common.by import By
from docx.shared import Inches
from selenium.webdriver.support import expected_conditions as EC
from datetime import datetime
from io import BytesIO
from docx.shared import Inches
from selenium.webdriver.common.by import By
import time

def check_move_activity_current_ame(driver, wait, document, current_ame_date):

    from datetime import datetime
    from io import BytesIO
    from docx.shared import Inches
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support import expected_conditions as EC
    import time

    document.add_heading("Daily Activity", level=2)

    try:
        ame_day = str(datetime.strptime(current_ame_date, "%m/%d/%Y").day)
    except:
        document.add_paragraph("Invalid AME date format.")
        return

    document.add_paragraph(
        f"Below are the move-in / move-out activities scheduled on AME date ({ame_day})."
    )

    # --------------------------------------------------
    # Helper to locate target cell fresh every time
    # --------------------------------------------------
    def get_target_cell():
        driver.switch_to.default_content()
        wait.until(EC.frame_to_be_available_and_switch_to_it((By.ID, "filter")))

        calendar_cells = driver.find_elements(
            By.XPATH,
            "//td[starts-with(@class,'Calendar')]"
        )

        for cell in calendar_cells:
            full_text = cell.text.strip()
            for token in full_text.split():
                if token.isdigit() and token == ame_day:
                    return cell
        return None

    # --------------------------------------------------
    # Process link safely
    # --------------------------------------------------
    def process_link(link_keyword):

        target_cell = get_target_cell()
        if not target_cell:
            return False

        try:
            link = target_cell.find_element(
                By.XPATH,
                f".//a[contains(text(),'{link_keyword}')]"
            )
        except:
            return False

        driver.execute_script("arguments[0].click();", link)

        move_table = wait.until(
            EC.presence_of_element_located(
                (By.XPATH, "//table[.//th[contains(text(),'Move')]]")
            )
        )

        time.sleep(1)

        png = move_table.screenshot_as_png
        image_stream = BytesIO(png)
        document.add_picture(image_stream, width=Inches(6))
        document.add_paragraph("")

        # Go back and re-enter iframe
        driver.switch_to.default_content()
        driver.back()
        wait.until(EC.frame_to_be_available_and_switch_to_it((By.ID, "filter")))
        time.sleep(1)

        return True

    found_move_in = process_link("Move in")
    found_move_out = process_link("Move out")

    if not found_move_in and not found_move_out:
        document.add_paragraph(
            f"No move-in or move-out scheduled on AME date ({ame_day})."
        )

    driver.switch_to.default_content()
