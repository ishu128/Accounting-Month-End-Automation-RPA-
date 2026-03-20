import os
import time
import pandas as pd
import matplotlib.pyplot as plt
import logging
log = logging.getLogger(__name__)
from datetime import datetime
from io import BytesIO
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support import expected_conditions as EC

from docx.shared import Inches

from utils.download_utils import wait_for_download_and_rename

def run_security_deposit_activity(driver, wait, prop_code, prop_name, folder_path):
    """
    Step 12 – Download Security Deposit Activity Report
    """

    driver.switch_to.default_content()
    log.info("🔍 Searching for Residential Analytics...")

    # Open search
    wait.until(EC.element_to_be_clickable((By.ID, "miSearch"))).click()

    search_box = wait.until(
        EC.visibility_of_element_located(
            (By.XPATH, "//li[@id='miSearch']//input[@type='text']")
        )
    )
    search_box.clear()
    search_box.send_keys("Residential Analytics")

    wait.until(
        EC.element_to_be_clickable(
            (By.XPATH, "//a[contains(text(),'Residential Analytics')]")
        )
    ).click()

    log.info("📊 Residential Analytics opened.")

    # Switch to correct iframe
    driver.switch_to.default_content()
    for frame in driver.find_elements(By.TAG_NAME, "iframe"):
        driver.switch_to.default_content()
        driver.switch_to.frame(frame)
        if driver.find_elements(By.ID, "ReportType_DropDownList"):
            break
    else:
        raise RuntimeError("❌ Residential Analytics iframe not found.")

    # Set property
    prop_box = wait.until(
        EC.presence_of_element_located((By.ID, "PropLookup_LookupCode"))
    )
    driver.execute_script("arguments[0].value = arguments[1];", prop_box, prop_code)
    prop_box.send_keys(Keys.TAB)

    # Select Security Deposit Activity
    Select(wait.until(
        EC.presence_of_element_located((By.ID, "ReportType_DropDownList"))
    )).select_by_visible_text("Security Deposit Activity")

    current_mmyy = datetime.today().strftime("%m/%Y")

    driver.execute_script(
        "arguments[0].value = arguments[1];",
        wait.until(EC.presence_of_element_located((By.ID, "MMYY1_TextBox"))),
        current_mmyy
    )

    driver.execute_script(
        "arguments[0].value = arguments[1];",
        wait.until(EC.presence_of_element_located((By.ID, "MMYY2_TextBox"))),
        current_mmyy
    )

    # Display
    wait.until(EC.element_to_be_clickable((By.ID, "Display_Button"))).click()
    time.sleep(4)

    # Download
    download_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    files_before = set(os.listdir(download_dir))

    wait.until(EC.element_to_be_clickable((By.ID, "Excel_Button"))).click()

    log.info("📥 Waiting for Excel download...")

    return wait_for_download_and_rename(
        download_dir=download_dir,
        files_before=files_before,
        target_folder=folder_path,
        base_filename=f"{prop_name}_Security_Deposit_Activity"
    )

import matplotlib.pyplot as plt
from io import BytesIO

def save_dataframe_snapshot(df):
    """
    Returns image stream of DataFrame snapshot.
    Nothing saved on disk.
    """

    if df.empty:
        return None

    df = df.copy()

    # Fix Unit formatting
    if "Unit" in df.columns:
        df["Unit"] = df["Unit"].apply(
            lambda x: f"Unit {int(x)}"
            if isinstance(x, (int, float)) and not pd.isna(x)
            else str(x)
        )

    fig, ax = plt.subplots(figsize=(12, 0.5 * len(df) + 1))

    ax.axis("off")

    table = ax.table(
        cellText=df.values,
        colLabels=df.columns,
        loc="center"
    )

    table.auto_set_font_size(False)
    table.set_fontsize(8)
    table.auto_set_column_width(col=list(range(len(df.columns))))

    plt.tight_layout()

    image_stream = BytesIO()
    plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=300)
    plt.close(fig)

    image_stream.seek(0)

    return image_stream

def review_security_deposit_report(file_path):

    df = pd.read_excel(file_path, skiprows=4)

    df.columns = [
        "Property", "Unit", "Resident Code", "Resident",
        "Prior Deposit Billed", "Prior Receipts",
        "Current Dep.Billed", "Current Receipts",
        "Deposits On Hand", "(Prpd)/Delnq Deposits",
        "Deposits Forfeited"
    ]

    # 🔥 Fix 1: Clean Resident column
    df["Resident"] = df["Resident"].fillna("").astype(str)

    # 🔥 Fix 2: Clean Deposit column
    df["Deposits On Hand"] = pd.to_numeric(
        df["Deposits On Hand"],
        errors="coerce"
    ).fillna(0)

    resident_lower = df["Resident"].str.lower()

    past_mask = resident_lower.str.contains(
        "past|canceled|denied",
        na=False   # 🔥 VERY IMPORTANT
    )

    past_df = df[
        past_mask &
        (df["Deposits On Hand"] != 0)
    ]

    negative_df = df[
        (~past_mask) &
        (df["Deposits On Hand"] < 0)
    ]

    return past_df, negative_df

    


from docx.shared import Inches
from docx import Document

from docx.shared import Inches

def add_sd_review_to_word(document, past_df, negative_df):

    document.add_heading(
        "Step 12 Review Security Deposit Activity",
        level=2
    )

    document.add_paragraph(
        "There are some negative amounts in Deposit on Hand or "
        "Past/Denied/Canceled residents with balance."
    )

    # Past / Canceled / Denied
    if not past_df.empty:

        document.add_heading(
            "Security Deposit – Past / Canceled / Denied With Balance",
            level=3
        )

        image_stream = save_dataframe_snapshot(past_df)

        if image_stream:
            document.add_picture(image_stream, width=Inches(6))


    # Negative
    if not negative_df.empty:

        document.add_heading(
            "Security Deposit – Negative Deposit On Hand",
            level=3
        )

        image_stream = save_dataframe_snapshot(negative_df)

        if image_stream:
            document.add_picture(image_stream, width=Inches(6))
            

def run_security_deposit_full(driver, wait, prop_code, prop_name, folder_path, document):

    sd_path = run_security_deposit_activity(
        driver, wait, prop_code, prop_name, folder_path
    )

    past_df, negative_df = review_security_deposit_report(sd_path)

    add_sd_review_to_word(
        document,
        past_df,
        negative_df,
    )

